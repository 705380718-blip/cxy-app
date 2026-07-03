from io import BytesIO
from pathlib import Path

from fastapi.testclient import TestClient
from docx import Document

from app.main import app


def get_project(client: TestClient, project_key: str) -> dict:
    projects = client.get("/projects")
    assert projects.status_code == 200
    return next(project for project in projects.json() if project["key"] == project_key)


def test_project_summary_and_dashboard_flag():
    client = TestClient(app)

    summary = client.get("/projects/summary")
    assert summary.status_code == 200
    assert summary.json()["total"] >= 4

    update = client.patch("/projects/ops/dashboard", json={"dashboard": True})
    assert update.status_code == 200
    assert update.json()["dashboard"] is True

    dashboard_projects = client.get("/projects", params={"dashboard": True})
    assert dashboard_projects.status_code == 200
    assert any(project["key"] == "ops" for project in dashboard_projects.json())


def test_create_project_and_task():
    client = TestClient(app)
    client.delete("/projects/unit-test-project")

    project = client.post(
        "/projects",
        json={
            "key": "unit-test-project",
            "name": "单元测试项目",
            "customer": "测试客户",
            "phase": "预投入",
            "manager": "测试经理",
            "dashboard": True,
        },
    )
    assert project.status_code == 200
    assert project.json()["key"] == "unit-test-project"

    task = client.post(
        "/tasks",
        json={
            "project_key": "unit-test-project",
            "status": "todo",
            "title": "准备测试任务",
            "description": "验证任务 API",
            "owner": "测试经理",
            "due_date": "2026-05-30",
        },
    )
    assert task.status_code == 200
    assert task.json()["project_key"] == "unit-test-project"

    tasks = client.get("/tasks", params={"project_key": "unit-test-project"})
    assert tasks.status_code == 200
    assert any(item["title"] == "准备测试任务" for item in tasks.json())


def test_project_gantt_rows_are_persisted():
    client = TestClient(app)

    rows = [
        {
            "label": "样例数据技术验证",
            "bars": [
                {
                    "text": "完成样例数据验证",
                    "startDate": "2026-06-01",
                    "endDate": "2026-06-05",
                    "status": "active",
                    "owner": "陈晓勇",
                    "note": "重启服务后仍应保留",
                }
            ],
        }
    ]

    saved = client.put("/projects/gov/gantt", json={"rows": rows})
    assert saved.status_code == 200
    assert saved.json()["rows"] == rows

    loaded = client.get("/projects/gov/gantt")
    assert loaded.status_code == 200
    assert loaded.json()["rows"] == rows


def test_task_status_calendar_and_risk_summary():
    client = TestClient(app)

    created = client.post(
        "/tasks",
        json={
            "project_key": "gov",
            "status": "todo",
            "title": "P2 单测任务",
            "owner": "测试经理",
            "due_date": "2026-05-27",
            "progress": 0,
        },
    )
    assert created.status_code == 200
    task_id = created.json()["id"]

    updated = client.patch(f"/tasks/{task_id}", json={"status": "doing", "progress": 50})
    assert updated.status_code == 200
    assert updated.json()["status"] == "doing"
    assert updated.json()["progress"] == 50

    calendar = client.get("/tasks/calendar", params={"project_key": "gov", "month": "2026-05"})
    assert calendar.status_code == 200
    assert any(item["id"] == task_id for item in calendar.json()["tasks"])
    assert any(item["title"] == "需求评审会" for item in calendar.json()["milestones"])

    summary = client.get("/risks/summary", params={"project_key": "gov"})
    assert summary.status_code == 200
    assert summary.json()["open_count"] >= 1
    assert summary.json()["avg_task_progress"] >= 0
    assert 0 <= summary.json()["health"] <= 100
    assert "overdue_tasks" in summary.json()


def test_milestone_status_can_be_updated():
    client = TestClient(app)

    created = client.post(
        "/milestones",
        json={
            "project_key": "gov",
            "title": "上线评审",
            "date": "2026-05-29",
            "status": "待开始",
        },
    )
    assert created.status_code == 200
    milestone_id = created.json()["id"]

    updated = client.patch(f"/milestones/{milestone_id}", json={"status": "进行中"})
    assert updated.status_code == 200
    assert updated.json()["status"] == "进行中"

    milestones = client.get("/milestones", params={"project_key": "gov"})
    assert milestones.status_code == 200
    assert any(item["id"] == milestone_id and item["status"] == "进行中" for item in milestones.json())


def test_milestone_can_be_deleted_from_calendar_and_list():
    client = TestClient(app)

    created = client.post(
        "/milestones",
        json={
            "project_key": "gov",
            "title": "待删除里程碑",
            "date": "2026-05-30",
            "status": "待开始",
        },
    )
    assert created.status_code == 200
    milestone_id = created.json()["id"]

    deleted = client.delete(f"/milestones/{milestone_id}")
    assert deleted.status_code == 200
    assert deleted.json() == {"status": "deleted", "id": milestone_id}

    milestones = client.get("/milestones", params={"project_key": "gov"})
    assert milestones.status_code == 200
    assert all(item["id"] != milestone_id for item in milestones.json())

    calendar = client.get("/tasks/calendar", params={"project_key": "gov", "month": "2026-05"})
    assert calendar.status_code == 200
    assert all(item["id"] != milestone_id for item in calendar.json()["milestones"])


def test_done_assessment_task_closes_linked_demand():
    client = TestClient(app)

    demand = client.post(
        "/demands",
        json={
            "project_key": "gov",
            "title": "增加评估联动单测需求",
            "description": "验证需求评估任务完成后联动需求状态",
            "status": "待评估",
            "scope_impact": "影响现有排期",
        },
    )
    assert demand.status_code == 200
    demand_id = demand.json()["id"]

    before = get_project(client, "gov")["health_breakdown"]
    assert before["signals"]["open_demands"] >= 1

    task = client.post(
        "/tasks",
        json={
            "project_key": "gov",
            "status": "todo",
            "title": "评估需求：增加评估联动单测需求",
            "description": "来自新增需求的评估任务",
            "owner": "测试经理",
            "due_date": "2026-05-30",
            "demand_id": demand_id,
        },
    )
    assert task.status_code == 200
    task_id = task.json()["id"]
    assert task.json()["demand_id"] == demand_id

    updated = client.patch(f"/tasks/{task_id}", json={"status": "done", "progress": 100})
    assert updated.status_code == 200
    assert updated.json()["status"] == "done"

    demands = client.get("/demands", params={"project_key": "gov"})
    assert demands.status_code == 200
    linked_demand = next(item for item in demands.json() if item["id"] == demand_id)
    assert linked_demand["status"] == "已完成"

    after = get_project(client, "gov")["health_breakdown"]
    assert after["signals"]["open_demands"] == before["signals"]["open_demands"] - 1
    assert after["health"] > before["health"]


def test_snippet_extract_edit_and_confirm():
    client = TestClient(app)

    snippet = client.post(
        "/snippets",
        json={
            "project_key": "gov",
            "source_type": "meeting",
            "raw_text": "5/28 完成天气 API 联调，责任人李工。风险：短信供应商文档未到位，可能影响评审。新增需求：首页增加办件进度提醒。",
        },
    )
    assert snippet.status_code == 200
    snippet_id = snippet.json()["id"]

    extracted = client.post(f"/snippets/{snippet_id}/extract")
    assert extracted.status_code == 200
    items = extracted.json()["items"]
    assert len(items) >= 3

    task_item = next(item for item in items if item["item_type"] == "task")
    updated = client.patch(
        f"/extractions/{task_item['id']}",
        json={"title": "完成天气 API 联调", "owner": "李工", "due_date": "2026-05-28"},
    )
    assert updated.status_code == 200
    assert updated.json()["title"] == "完成天气 API 联调"

    confirmed = client.post(f"/extractions/{task_item['id']}/confirm")
    assert confirmed.status_code == 200
    assert confirmed.json()["status"] == "confirmed"
    assert confirmed.json()["target_table"] == "tasks"

    tasks = client.get("/tasks", params={"project_key": "gov"})
    assert any(task["source_extraction_id"] == task_item["id"] for task in tasks.json())


def test_document_template_generate_and_export_word():
    client = TestClient(app)

    upload = client.post(
        "/document-templates/upload",
        data={"name": "单元测试 SRS 模板", "template_type": "srs"},
        files={
            "file": (
                "unit-template.md",
                b"# template\n\nProject: {{ project.name }}",
                "text/markdown",
            )
        },
    )
    assert upload.status_code == 200
    template = upload.json()
    assert template["name"] == "单元测试 SRS 模板"
    assert template["original_filename"] == "unit-template.md"

    generated = client.post(
        "/document-versions/generate",
        json={
            "project_key": "gov",
            "template_id": template["id"],
            "title": "单元测试需求规格说明书",
        },
    )
    assert generated.status_code == 200
    version = generated.json()
    assert "项目背景" in version["content_markdown"]
    assert version["word_status"] == "not_exported"

    exported = client.post(f"/document-versions/{version['id']}/export-word")
    assert exported.status_code == 200
    exported_version = exported.json()
    assert exported_version["word_status"] == "exported"
    assert exported_version["export_url"].endswith(".docx")

    download = client.get(exported_version["export_url"].replace("/api", ""))
    assert download.status_code == 200
    export_path = exported_version["export_path"]
    document = Document(export_path)
    assert any("单元测试需求规格说明书" in paragraph.text for paragraph in document.paragraphs)


def test_delete_document_template_removes_template_file_and_versions():
    client = TestClient(app)

    upload = client.post(
        "/document-templates/upload",
        data={"name": "待删除模板", "template_type": "weekly"},
        files={
            "file": (
                "delete-template.md",
                b"# delete template",
                "text/markdown",
            )
        },
    )
    assert upload.status_code == 200
    template = upload.json()
    template_path = Path(template["file_path"])
    assert template_path.exists()

    generated = client.post(
        "/document-versions/generate",
        json={
            "project_key": "gov",
            "template_id": template["id"],
            "title": "待删除模板版本",
        },
    )
    assert generated.status_code == 200
    version_id = generated.json()["id"]

    deleted = client.delete(f"/document-templates/{template['id']}")
    assert deleted.status_code == 200
    assert deleted.json()["deleted"] is True
    assert not template_path.exists()

    templates = client.get("/document-templates")
    assert all(item["id"] != template["id"] for item in templates.json())

    versions = client.get("/document-versions", params={"template_id": template["id"]})
    assert all(item["id"] != version_id for item in versions.json())


def docx_bytes(document: Document) -> bytes:
    stream = BytesIO()
    document.save(stream)
    stream.seek(0)
    return stream.read()


def test_docx_content_extraction_template_mapping_and_table_export():
    client = TestClient(app)

    template_doc = Document()
    template_doc.add_heading("引言", level=1)
    template_doc.add_heading("目的", level=2)
    template_doc.add_heading("测试进度", level=1)
    template_doc.add_heading("整体进度", level=2)
    template_doc.add_heading("实施任务和时间人员安排", level=2)
    upload = client.post(
        "/document-templates/upload",
        data={"name": "测试方案 Word 模板", "template_type": "test_plan"},
        files={
            "file": (
                "test-plan-template.docx",
                docx_bytes(template_doc),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert upload.status_code == 200
    template = upload.json()

    content_doc = Document()
    content_doc.add_heading("文档目的", level=1)
    content_doc.add_paragraph("明确测试目标和执行范围。")
    content_doc.add_heading("测试阶段安排", level=1)
    table = content_doc.add_table(rows=2, cols=4)
    table.rows[0].cells[0].text = "测试阶段"
    table.rows[0].cells[1].text = "主要任务和工作"
    table.rows[0].cells[2].text = "准入/准出依据"
    table.rows[0].cells[3].text = "输出工作产品"
    table.rows[1].cells[0].text = "阶段一\n2026.06.01-2026.06.07"
    table.rows[1].cells[1].text = "需求确认\n测试用例编写"
    table.rows[1].cells[2].text = "需求冻结\n用例评审通过"
    table.rows[1].cells[3].text = "测试计划\n测试用例"
    extracted = client.post(
        "/document-content/extract",
        files={
            "file": (
                "test-content.docx",
                docx_bytes(content_doc),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert extracted.status_code == 200
    extracted_content = extracted.json()["content"]
    assert "阶段一<br>2026.06.01-2026.06.07" in extracted_content

    generated = client.post(
        "/document-versions/generate",
        json={
            "project_key": "gov",
            "template_id": template["id"],
            "title": "测试方案映射验证",
            "input_content": extracted_content,
        },
    )
    assert generated.status_code == 200
    version = generated.json()
    assert version["content_markdown"].count("阶段一<br>2026.06.01-2026.06.07") == 1
    assert "## 测试进度\n\n待补充" not in version["content_markdown"]
    assert "待补充。" not in version["content_markdown"]

    exported = client.post(f"/document-versions/{version['id']}/export-word")
    assert exported.status_code == 200
    document = Document(exported.json()["export_path"])
    assert len(document.tables) == 2
    assert document.tables[0].autofit is False
    assert "阶段一\n2026.06.01-2026.06.07" in document.tables[0].cell(1, 0).text
    assert "角色" in document.tables[1].cell(0, 0).text


def test_document_generation_quality_rules_ground_project_and_reduce_duplicates():
    client = TestClient(app)

    template_doc = Document()
    for level, heading in [
        (1, "引言"),
        (2, "目的"),
        (2, "背景"),
        (2, "范围"),
        (2, "测试参考文档"),
        (2, "测试提交文档"),
        (1, "测试进度"),
        (2, "整体进度"),
        (2, "实施任务和时间人员安排"),
        (1, "环境要求"),
        (2, "硬件环境"),
        (2, "软件测试环境"),
        (2, "测试工具"),
        (1, "测试方法"),
        (2, "功能测试"),
        (2, "接口测试"),
        (2, "性能测试"),
        (1, "质量目标"),
        (1, "缺陷严重度描述"),
    ]:
        template_doc.add_heading(heading, level=level)
    upload = client.post(
        "/document-templates/upload",
        data={"name": "高优先级质量规则模板", "template_type": "test_plan"},
        files={
            "file": (
                "quality-template.docx",
                docx_bytes(template_doc),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert upload.status_code == 200

    content = """
# 文档目的
本文档以轩辕大数据平台的需求和界面原型为依据，对系统测试需求进行功能点分析。

# 项目概况
项目名称：轩辕大数据平台建设项目
实施地点: 成都云计算中心办公楼

# 测试阶段安排
| 测试阶段 | 主要任务和工作 | 准入/准出依据 | 输出工作产品 |
| --- | --- | --- | --- |
| 测试准备（第1周） 2024.3.25~2024.3.31 | 编写测试计划 | 需求确认 | 测试计划 |
"""
    generated = client.post(
        "/document-versions/generate",
        json={
            "project_key": "gov",
            "template_id": upload.json()["id"],
            "title": "智慧政务平台测试方案",
            "input_content": content,
        },
    )
    assert generated.status_code == 200
    markdown = generated.json()["content_markdown"]
    assert "轩辕大数据平台建设项目" not in markdown
    assert "成都云计算中心办公楼" not in markdown
    assert "2024.3.25" not in markdown
    assert "待补充。" not in markdown
    assert markdown.count("| 测试阶段 |") == 1
    assert "| 角色 | 人员/团队 | 主要职责 | 备注 |" in markdown
    assert "| 环境项 | 建议配置 | 说明 |" in markdown
    assert "| 软件项 | 版本/要求 | 说明 |" in markdown
    assert "| 严重度 | 定义 | 处理要求 |" in markdown


def test_word_export_falls_back_when_template_has_no_bullet_style():
    client = TestClient(app)

    template_doc = Document()
    template_doc.add_heading("测试内容", level=1)
    upload = client.post(
        "/document-templates/upload",
        data={"name": "无项目符号样式模板", "template_type": "weekly"},
        files={
            "file": (
                "no-bullet-style-template.docx",
                docx_bytes(template_doc),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert upload.status_code == 200

    generated = client.post(
        "/document-versions/generate",
        json={
            "project_key": "gov",
            "template_id": upload.json()["id"],
            "title": "项目符号降级验证",
            "input_content": "# 测试内容\n- 第一条\n- 第二条",
        },
    )
    assert generated.status_code == 200

    exported = client.post(f"/document-versions/{generated.json()['id']}/export-word")
    assert exported.status_code == 200
    document = Document(exported.json()["export_path"])
    assert any("第一条" in paragraph.text for paragraph in document.paragraphs)


def test_word_export_preserves_docx_template_tables_and_headings():
    client = TestClient(app)

    template_doc = Document()
    template_doc.add_paragraph("05 系统测试方案_XXX系统Vx.x")
    cover = template_doc.add_table(rows=2, cols=2)
    cover.cell(0, 0).text = "编制单位"
    cover.cell(0, 1).text = "Xxxx"
    cover.cell(1, 0).text = "编制日期"
    cover.cell(1, 1).text = ""
    template_doc.add_heading("1 引言", level=1)
    template_doc.add_heading("1.1 目的", level=2)
    template_doc.add_paragraph("本文档说明 XXX系统 的测试目的。")
    template_doc.add_heading("1.2 范围", level=2)
    scope = template_doc.add_table(rows=4, cols=2)
    scope.cell(0, 0).text = "模块名称"
    scope.cell(0, 1).text = "功能名称"
    template_doc.add_heading("1.3 测试参考文档", level=2)
    refs = template_doc.add_table(rows=10, cols=5)
    for index, value in enumerate(["文档 / （版本/日期）", "已创建或可用", "已被接收或已经过复审", "作者或来源", "备注"]):
        refs.cell(0, index).text = value
    template_doc.add_heading("2 测试进度", level=1)
    schedule = template_doc.add_table(rows=6, cols=4)
    for index, value in enumerate(["测试阶段", "开始时间", "完成时间", "测试人员"]):
        schedule.cell(0, index).text = value
    template_doc.add_heading("3 测试方法", level=1)
    template_doc.add_heading("3.1 功能测试", level=2)
    method = template_doc.add_table(rows=7, cols=2)
    for index, value in enumerate(["测试目标", "测试范围：", "开始标准：", "完成标准：", "测试重点和优先级：", "需考虑的特殊事项：", "备注："]):
        method.cell(index, 0).text = value
        method.cell(index, 1).text = "模板占位"

    upload = client.post(
        "/document-templates/upload",
        data={"name": "模板保留测试方案", "template_type": "test_plan"},
        files={
            "file": (
                "preserve-template.docx",
                docx_bytes(template_doc),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert upload.status_code == 200

    generated = client.post(
        "/document-versions/generate",
        json={
            "project_key": "gov",
            "template_id": upload.json()["id"],
            "title": "模板保留验证",
        },
    )
    assert generated.status_code == 200

    exported = client.post(f"/document-versions/{generated.json()['id']}/export-word")
    assert exported.status_code == 200
    document = Document(exported.json()["export_path"])
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "XXX系统" not in paragraph_text
    assert "智慧政务平台建设项目" in paragraph_text
    assert len(document.tables) == len(template_doc.tables)
    assert document.tables[1].cell(1, 0).text.strip()
    assert document.tables[2].cell(1, 0).text == "需求规格说明书/需求清单"
    assert document.tables[3].cell(1, 0).text == "测试准备"
    assert document.tables[4].cell(0, 0).text == "测试目标"
    assert document.tables[4].cell(6, 0).text == "备注："
    heading_styles = [(paragraph.style.name, paragraph.text) for paragraph in document.paragraphs if paragraph.text.strip()]
    assert ("Heading 1", "1 引言") in heading_styles
    assert ("Heading 2", "1.1 目的") in heading_styles


def test_word_export_maps_source_tables_into_docx_template_tables():
    client = TestClient(app)

    template_doc = Document()
    template_doc.add_heading("测试进度", level=1)
    template_doc.add_heading("整体进度", level=2)
    schedule = template_doc.add_table(rows=4, cols=4)
    for index, value in enumerate(["测试阶段", "开始时间", "完成时间", "测试人员"]):
        schedule.cell(0, index).text = value
    template_doc.add_heading("环境要求", level=1)
    template_doc.add_heading("软件测试环境", level=2)
    software = template_doc.add_table(rows=3, cols=2)
    software.cell(0, 0).text = "软件需求"
    software.cell(0, 1).text = "用途"

    upload = client.post(
        "/document-templates/upload",
        data={"name": "表格映射模板", "template_type": "test_plan"},
        files={
            "file": (
                "source-table-template.docx",
                docx_bytes(template_doc),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert upload.status_code == 200

    content = """
# 测试阶段安排
| 测试阶段 | 主要任务和工作 | 准入/准出依据 | 输出工作产品 |
| --- | --- | --- | --- |
| 阶段一 2026.06.01-2026.06.07 | 需求确认、测试用例编写 | 需求冻结 | 测试计划、测试用例 |

# 软件测试环境
| 软件需求 | 用途 |
| --- | --- |
| MySQL 8.0 | 测试数据存储 |
"""

    generated = client.post(
        "/document-versions/generate",
        json={
            "project_key": "gov",
            "template_id": upload.json()["id"],
            "title": "表格映射验证",
            "input_content": content,
        },
    )
    assert generated.status_code == 200

    exported = client.post(f"/document-versions/{generated.json()['id']}/export-word")
    assert exported.status_code == 200
    document = Document(exported.json()["export_path"])
    assert document.tables[0].cell(1, 0).text == "阶段一"
    assert document.tables[0].cell(1, 1).text == "2026-06-01"
    assert document.tables[0].cell(1, 2).text == "2026-06-07"
    assert document.tables[1].cell(1, 0).text == "MySQL 8.0"
    assert document.tables[1].cell(1, 1).text == "测试数据存储"


def test_model_config_save_test_and_assistant_context_answer():
    client = TestClient(app)

    config = client.get("/settings/model-config")
    assert config.status_code == 200
    payload = config.json()
    payload["provider"] = "mock"
    payload["provider_label"] = "离线内置模型"
    payload["model"] = "local-project-assistant"
    payload["api_key"] = "secret-for-test"

    saved = client.put("/settings/model-config", json=payload)
    assert saved.status_code == 200
    assert saved.json()["model"] == "local-project-assistant"
    assert saved.json()["api_key"] == "********"
    assert saved.json()["has_api_key"] is True

    tested = client.post("/settings/model-config/test", json=saved.json())
    assert tested.status_code == 200
    assert tested.json()["status"] == "connected"
    assert tested.json()["config"]["api_key"] == "********"

    answer = client.post(
        "/assistant/chat",
        json={
            "project_key": "gov",
            "view": "项目驾驶舱",
            "message": "这个项目现在最大风险是什么？",
        },
    )
    assert answer.status_code == 200
    assert answer.json()["source"] == "local-project-context"
    assert "风险" in answer.json()["answer"] or "关注" in answer.json()["answer"]

    task_answer = client.post(
        "/assistant/chat",
        json={
            "project_key": "gov",
            "view": "模型与 Agent",
            "message": "创建一个接口文档任务",
        },
    )
    assert task_answer.status_code == 200
    task_payload = task_answer.json()
    assert task_payload["action_type"] == "pending_extraction"
    assert task_payload["action_payload"]["item_type"] == "task"

    extraction_id = task_payload["action_payload"]["extraction_id"]
    extractions = client.get("/extractions", params={"project_key": "gov", "status": "pending"})
    assert any(item["id"] == extraction_id for item in extractions.json())

    messages = client.get("/assistant/messages", params={"project_key": "gov"})
    assert messages.status_code == 200
    assert any(message["content"] == "创建一个接口文档任务" for message in messages.json())

    domestic_payload = saved.json()
    domestic_payload["provider"] = "deepseek"
    domestic_payload["provider_label"] = "DeepSeek"
    domestic_payload["model"] = "deepseek-chat"
    domestic_payload["base_url"] = "https://api.deepseek.com"
    domestic_payload["api_key"] = "deepseek-secret"

    domestic_saved = client.put("/settings/model-config", json=domestic_payload)
    assert domestic_saved.status_code == 200
    domestic_config = domestic_saved.json()
    assert domestic_config["provider"] == "deepseek"
    assert domestic_config["provider_label"] == "DeepSeek"
    assert domestic_config["base_url"] == "https://api.deepseek.com"
    assert domestic_config["api_key"] == "********"

    lm_studio_payload = domestic_config
    lm_studio_payload["provider"] = "lm-studio"
    lm_studio_payload["provider_label"] = "本地 LM Studio"
    lm_studio_payload["model"] = "local-model"
    lm_studio_payload["base_url"] = "http://localhost:1234/v1"
    lm_studio_payload["api_key"] = ""

    lm_studio_saved = client.put("/settings/model-config", json=lm_studio_payload)
    assert lm_studio_saved.status_code == 200
    lm_studio_config = lm_studio_saved.json()
    assert lm_studio_config["provider"] == "lm-studio"
    assert lm_studio_config["provider_label"] == "本地 LM Studio"
    assert lm_studio_config["base_url"] == "http://localhost:1234/v1"

    restore_payload = lm_studio_config
    restore_payload["provider"] = "mock"
    restore_payload["provider_label"] = "离线内置模型"
    restore_payload["model"] = "local-project-assistant"
    restore_payload["base_url"] = ""
    restore_payload["api_key"] = ""
    restored = client.put("/settings/model-config", json=restore_payload)
    assert restored.status_code == 200
