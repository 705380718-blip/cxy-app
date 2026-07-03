from __future__ import annotations

from datetime import date, datetime, timedelta
from html import escape
from pathlib import Path
import re

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


def rows(connection, sql: str, params: list | tuple) -> list[dict]:
    return [dict(row) for row in connection.execute(sql, params).fetchall()]


def get_project_context(connection, project_key: str) -> dict:
    project = connection.execute("select * from projects where key = ?", [project_key]).fetchone()
    if not project:
        raise ValueError("Project not found")
    return {
        "project": dict(project),
        "tasks": rows(connection, "select * from tasks where project_key = ? order by due_date asc, id desc", [project_key]),
        "risks": rows(connection, "select * from risks where project_key = ? order by id desc", [project_key]),
        "demands": rows(connection, "select * from demands where project_key = ? order by id desc", [project_key]),
        "milestones": rows(connection, "select * from milestones where project_key = ? order by date asc, id desc", [project_key]),
    }


def bullet_list(items: list[dict], fields: tuple[str, ...], empty_text: str) -> str:
    if not items:
        return f"- {empty_text}"
    lines = []
    for item in items[:12]:
        parts = [str(item.get(field) or "") for field in fields if item.get(field)]
        lines.append(f"- {' | '.join(parts)}")
    return "\n".join(lines)


def parse_date(value: str) -> date | None:
    try:
        return datetime.strptime(str(value or "")[:10], "%Y-%m-%d").date()
    except ValueError:
        return None


def project_period(context: dict) -> tuple[date | None, date | None]:
    project = context["project"]
    start = parse_date(project.get("start_date") or project.get("pre_start_date") or "")
    end = parse_date(project.get("end_date") or "")
    return start, end


def date_range_text(start: date | None, days: int, fallback: str) -> str:
    if not start:
        return fallback
    end = start + timedelta(days=max(days - 1, 0))
    return f"{start.isoformat()}~{end.isoformat()}"


def project_overview_lines(context: dict) -> list[str]:
    project = context["project"]
    lines = [
        f"- 项目名称：{project.get('name') or '项目名称待确认'}",
        f"- 客户/业主单位：{project.get('customer') or '客户单位待确认'}",
        f"- 当前阶段：{project.get('phase') or '阶段待确认'}",
    ]
    if project.get("region"):
        lines.append(f"- 项目地区：{project['region']}")
    if project.get("manager"):
        lines.append(f"- 项目经理：{project['manager']}")
    return lines


def project_scope_text(context: dict) -> str:
    project = context["project"]
    plan = project.get("plan") or "以项目计划、需求文档和评审结论为准。"
    return "\n".join(
        [
            f"本次文档适用于 {project.get('name') or '当前项目'} 的交付测试与质量确认工作。",
            f"测试范围围绕项目计划展开：{plan}",
            "不在本轮范围内的内容，应在评审阶段形成边界说明并纳入后续变更管理。",
        ]
    )


def test_schedule_table(context: dict) -> str:
    start, _ = project_period(context)
    rows = [
        ("测试准备", date_range_text(start, 3, "以项目启动和需求冻结日期为准"), "需求确认、测试计划编制、测试环境检查", "需求范围已确认；环境和账号可用", "测试计划、测试用例初稿"),
        (
            "接口测试",
            date_range_text(start + timedelta(days=3) if start else None, 5, "以接口联调窗口为准"),
            "接口连通性、异常分支、权限与数据口径验证",
            "主要接口已部署并具备测试数据",
            "接口测试记录、缺陷清单",
        ),
        (
            "系统测试",
            date_range_text(start + timedelta(days=8) if start else None, 7, "以系统测试窗口为准"),
            "核心业务流程、权限、报表、数据一致性、易用性测试",
            "版本准入通过；关键缺陷完成修复",
            "系统测试用例执行记录、缺陷闭环记录",
        ),
        (
            "测试总结",
            date_range_text(start + timedelta(days=15) if start else None, 2, "以验收前总结节点为准"),
            "缺陷复盘、遗留问题确认、质量结论输出",
            "阻塞问题已处理或形成遗留确认",
            "测试总结报告、遗留问题清单",
        ),
    ]
    lines = [
        "| 测试阶段 | 时间安排 | 主要任务和工作 | 准入/准出依据 | 输出工作产品 |",
        "| --- | --- | --- | --- | --- |",
    ]
    for row in rows:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def test_personnel_table(context: dict) -> str:
    project = context["project"]
    manager = project.get("manager") or "项目经理待确认"
    delivery = project.get("delivery") or "交付负责人待确认"
    return "\n".join(
        [
            "| 角色 | 人员/团队 | 主要职责 | 备注 |",
            "| --- | --- | --- | --- |",
            f"| 项目经理 | {manager} | 统筹测试计划、客户沟通、风险协调和验收推进 | 对测试结论负责 |",
            f"| 交付/实施负责人 | {delivery} | 配合环境准备、版本部署、问题定位和修复验证 | 与测试人员保持日常同步 |",
            "| 测试负责人 | 待确认 | 组织测试设计、执行跟踪、缺陷复盘和测试报告输出 | 生成文档后需明确到人 |",
            "| 客户/业务代表 | 待确认 | 确认业务口径、验收标准和遗留问题处理意见 | 参与关键评审节点 |",
        ]
    )


def environment_table(context: dict, kind: str = "all") -> str:
    project = context["project"]
    region = project.get("region") or "项目现场/客户指定环境"
    if kind == "hardware":
        return "\n".join(
            [
                "| 环境项 | 建议配置 | 说明 |",
                "| --- | --- | --- |",
                "| 应用服务器 | 以部署方案或现场资源为准 | 用于部署被测系统服务 |",
                "| 数据库服务器 | 以数据库设计和实际数据量评估 | 需提前准备备份与恢复方案 |",
                "| 网络与访问 | 客户内网/VPN/白名单策略 | 确保测试账号可访问相关服务 |",
                f"| 部署地点 | {region} | 与项目实施地点保持一致 |",
            ]
        )
    if kind == "software":
        return "\n".join(
            [
                "| 软件项 | 版本/要求 | 说明 |",
                "| --- | --- | --- |",
                "| 操作系统 | 以部署方案为准 | 记录实际测试环境版本 |",
                "| 数据库 | 以数据库设计文档为准 | 需明确字符集、备份策略和账号权限 |",
                "| Web/应用中间件 | 以部署实施文档为准 | 需记录端口、域名和证书配置 |",
                "| 浏览器/客户端 | 覆盖客户常用终端 | 用于兼容性和易用性验证 |",
            ]
        )
    return "\n".join(
        [
            "| 类别 | 检查项 | 要求 |",
            "| --- | --- | --- |",
            "| 硬件环境 | 服务器、存储、网络 | 满足系统部署和并发测试需要 |",
            "| 软件环境 | 操作系统、数据库、中间件 | 与部署方案保持一致 |",
            "| 账号权限 | 测试账号、管理账号、接口账号 | 权限边界清晰，可追踪操作记录 |",
            "| 测试数据 | 基础数据、业务样例、异常数据 | 数据来源合法，敏感信息需脱敏 |",
        ]
    )


def test_tools_table() -> str:
    return "\n".join(
        [
            "| 序号 | 工具类型 | 工具名称 | 用途 |",
            "| --- | --- | --- | --- |",
            "| 1 | 接口测试工具 | Postman/JMeter | 接口连通性、参数校验和批量请求验证 |",
            "| 2 | 性能测试工具 | JMeter | 并发、响应时间和稳定性验证 |",
            "| 3 | 缺陷管理工具 | 项目缺陷台账/禅道/飞书表格 | 缺陷登记、流转和闭环跟踪 |",
            "| 4 | 日志与监控工具 | 系统日志/服务器监控 | 问题定位和资源占用分析 |",
        ]
    )


def test_method_text(title: str, context: dict) -> str:
    title_text = title.replace("测试", "")
    focus = {
        "功能": "围绕核心业务流程、权限控制、数据录入、查询统计和异常分支开展测试。",
        "接口": "覆盖接口鉴权、必填参数、异常返回、数据一致性和幂等性场景。",
        "集成": "验证系统模块之间、与第三方系统之间的数据流转和状态同步。",
        "用户界面": "检查页面布局、表单提示、操作反馈、兼容性和关键路径易用性。",
        "性能": "关注关键接口响应时间、并发处理能力、资源占用和长时间运行稳定性。",
        "回归": "围绕已修复缺陷、核心流程和高风险模块进行复测。",
        "负载": "通过逐步增加并发和数据量，观察系统容量边界和退化表现。",
        "安全": "检查账号权限、越权访问、弱口令、敏感数据展示和基础安全配置。",
    }
    content = next((value for key, value in focus.items() if key in title), "")
    return "\n".join(
        [
            content or f"围绕{title_text or '相关'}场景制定测试用例并执行验证。",
            "测试执行过程中应记录测试数据、操作步骤、实际结果和缺陷编号。",
            "通过标准为关键流程通过、阻塞缺陷关闭、遗留问题经项目相关方确认。",
        ]
    )


def reference_docs() -> str:
    return "\n".join(
        [
            "| 文档名称 | 用途 | 状态 |",
            "| --- | --- | --- |",
            "| 需求规格说明书/需求清单 | 明确测试范围和验收口径 | 待项目组确认最新版 |",
            "| 原型图/界面设计稿 | 支撑页面功能和交互验证 | 待项目组确认最新版 |",
            "| 接口设计文档 | 支撑接口测试和联调验证 | 待项目组确认最新版 |",
            "| 部署实施文档 | 支撑环境准备和版本部署检查 | 待项目组确认最新版 |",
            "| 历史缺陷及变更记录 | 支撑回归测试范围确认 | 持续更新 |",
        ]
    )


def submission_docs() -> str:
    return "\n".join(
        [
            "| 提交物 | 说明 | 责任方 |",
            "| --- | --- | --- |",
            "| 测试计划 | 测试范围、进度、人员和准入准出标准 | 测试负责人 |",
            "| 测试用例 | 各模块测试场景、步骤和预期结果 | 测试负责人 |",
            "| 缺陷清单 | 缺陷描述、等级、责任人、状态和闭环记录 | 测试/开发负责人 |",
            "| 测试报告 | 测试执行情况、质量结论和遗留问题 | 测试负责人 |",
            "| 验收问题清单 | 需客户或项目组确认的问题和处理意见 | 项目经理 |",
        ]
    )


def quality_goals() -> str:
    return "\n".join(
        [
            "| 指标 | 目标 | 说明 |",
            "| --- | --- | --- |",
            "| 核心用例通过率 | 100% | 核心业务流程必须全部通过 |",
            "| 总体用例通过率 | 不低于 95% | 未通过用例需说明原因和处理计划 |",
            "| 阻塞/严重缺陷 | 0 个未关闭 | 上线或验收前必须闭环 |",
            "| 一般缺陷 | 经评审可遗留 | 需明确责任人和计划完成时间 |",
            "| 遗留问题 | 项目相关方确认 | 形成书面记录并纳入后续跟踪 |",
        ]
    )


def defect_severity_table() -> str:
    return "\n".join(
        [
            "| 严重度 | 定义 | 处理要求 |",
            "| --- | --- | --- |",
            "| 致命 | 系统不可用、核心流程完全阻断或造成严重数据错误 | 立即处理，修复后优先回归 |",
            "| 严重 | 重要功能不可用、权限异常或影响关键业务办理 | 高优先级处理，纳入版本准出检查 |",
            "| 一般 | 局部功能异常、提示不准确或影响操作效率 | 按计划修复并回归验证 |",
            "| 轻微 | 文案、样式、兼容性等轻微问题 | 可结合版本节奏处理 |",
        ]
    )


TEMPLATE_TYPE_LABELS = {
    "srs": "需求规格说明书",
    "prd": "产品需求文档",
    "technical_solution": "技术方案",
    "api_design": "接口设计文档",
    "database_design": "数据库设计文档",
    "test_plan": "测试方案",
    "deployment_plan": "部署实施方案",
    "operation_manual": "运维手册",
    "user_manual": "用户使用手册",
    "acceptance_report": "验收报告",
    "weekly": "项目周报",
    "meeting_minutes": "会议纪要",
}


def user_input_section(input_content: str) -> str:
    content = input_content.strip()
    if not content:
        return "暂无新增输入内容。本稿主要根据当前项目库上下文生成。"
    return content


def markdown_heading_level(line: str) -> int:
    match = re.match(r"^(#{1,6})\s+", line)
    return len(match.group(1)) if match else 0


def normalize_heading(text: str) -> str:
    return re.sub(r"[^0-9A-Za-z\u4e00-\u9fa5]+", "", text).lower()


def strip_heading_number(text: str) -> str:
    return re.sub(r"^\s*\d+(?:\.\d+)*\s*", "", text).strip()


def parse_markdown_sections(markdown: str) -> dict[str, str]:
    sections: dict[str, list[str]] = {}
    current_title = ""
    current_lines: list[str] = []

    def flush() -> None:
        if current_title and current_lines:
            key = normalize_heading(strip_heading_number(current_title))
            sections.setdefault(key, []).extend(current_lines)

    for line in markdown.splitlines():
        level = markdown_heading_level(line)
        if level:
            flush()
            current_title = line[level + 1 :].strip()
            current_lines = []
            continue
        if current_title:
            current_lines.append(line)
    flush()
    return {key: "\n".join(lines).strip() for key, lines in sections.items() if "\n".join(lines).strip()}


def section_content(sections: dict[str, str], *candidates: str, used_keys: set[str] | None = None) -> tuple[str, str]:
    for candidate in candidates:
        key = normalize_heading(candidate)
        if used_keys is not None and key in used_keys:
            continue
        value = sections.get(key, "").strip()
        if value:
            return value, key
    return "", ""


def template_headings(template_path: str) -> list[tuple[int, str]]:
    if not template_path or Path(template_path).suffix.lower() != ".docx" or not Path(template_path).exists():
        return []
    document = Document(template_path)
    headings: list[tuple[int, str]] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        match = re.search(r"Heading\s+(\d+)|标题\s*(\d+)", style_name, re.I)
        if match:
            level = int(match.group(1) or match.group(2) or 1)
            headings.append((min(level, 3), text))
    return headings


def ground_project_content(content: str, context: dict, template_type: str) -> str:
    project = context["project"]
    lines = []
    for raw_line in content.splitlines():
        line = raw_line.strip()
        if re.match(r"^-?\s*项目名称[:：]", line):
            lines.append(f"项目名称：{project.get('name') or '项目名称待确认'}")
        elif re.match(r"^-?\s*(客户|业主单位|客户/业主单位)[:：]", line):
            lines.append(f"客户/业主单位：{project.get('customer') or '客户单位待确认'}")
        elif re.match(r"^-?\s*(实施地点|项目地点|项目地区)[:：]", line):
            lines.append(f"项目地区：{project.get('region') or '项目地区待确认'}")
        elif re.match(r"^-?\s*当前阶段[:：]", line):
            lines.append(f"当前阶段：{project.get('phase') or '阶段待确认'}")
        elif "本文档以" in line and "需求" in line and project.get("name") and project["name"] not in line:
            lines.append(f"本文档以 {project['name']} 的需求、项目资料和评审结论为依据，明确测试范围、测试方法和质量标准，指导测试用例编写、测试执行和验收判断。")
        else:
            lines.append(raw_line)
    grounded = "\n".join(lines).strip()
    if template_type in {"test_plan", "acceptance_report"} and "2024" in grounded:
        start, end = project_period(context)
        if start and start.year != 2024:
            grounded = re.sub(r"\b2024[./-]\d{1,2}[./-]\d{1,2}\s*[~-]\s*2024[./-]\d{1,2}[./-]\d{1,2}\b", "以项目实际测试计划为准", grounded)
            grounded = re.sub(r"\b2024[./-]\d{1,2}[./-]\d{1,2}\b", start.isoformat(), grounded)
        elif end and end.year != 2024:
            grounded = re.sub(r"\b2024[./-]\d{1,2}[./-]\d{1,2}\b", "项目计划日期", grounded)
    return grounded


def fallback_section_content(title: str, context: dict, template_type: str, used_keys: set[str]) -> str:
    if "参考文档" in title:
        return reference_docs()
    if "提交文档" in title or "交付文档" in title:
        return submission_docs()
    if "实施任务" in title or "人员安排" in title:
        return test_personnel_table(context)
    if "整体进度" in title or ("测试进度" in title and "测试阶段安排" not in used_keys):
        return test_schedule_table(context)
    if "硬件环境" in title:
        return environment_table(context, "hardware")
    if "软件测试环境" in title or "软件环境" in title:
        return environment_table(context, "software")
    if "测试工具" in title:
        return test_tools_table()
    if any(keyword in title for keyword in ("功能测试", "接口测试", "集成测试", "用户界面测试", "性能测试", "回归测试", "负载测试", "安全测试")):
        return test_method_text(title, context)
    if "质量目标" in title:
        return quality_goals()
    if "缺陷严重度" in title or "缺陷等级" in title:
        return defect_severity_table()
    if "目的" in title:
        return f"本文档用于明确 {context['project'].get('name') or '当前项目'} 的测试目标、范围、计划、方法、准入准出标准和风险应对措施，为测试执行、缺陷跟踪和验收决策提供依据。"
    if "背景" in title:
        project = context["project"]
        return "\n".join(project_overview_lines(context) + [project.get("background") or "项目背景以立项资料、合同范围和需求评审结论为准。"])
    if "范围" in title and "测试" not in title:
        return project_scope_text(context)
    if "风险" in title:
        return bullet_list(context["risks"], ("title", "impact", "status", "response"), "当前暂无已登记风险；测试过程中发现的环境、进度、接口和数据问题需及时纳入风险台账。")
    if template_type in {"test_plan", "acceptance_report"}:
        return "本节需结合项目最新资料持续维护，生成后由项目经理、测试负责人和相关干系人共同确认。"
    return "本节内容将结合项目上下文、用户输入资料和评审结论持续完善。"


def mapped_template_section(title: str, sections: dict[str, str], context: dict, used_keys: set[str]) -> tuple[str, str]:
    normalized = normalize_heading(strip_heading_number(title))
    direct = sections.get(normalized, "") if normalized not in used_keys else ""
    if direct:
        return direct, normalized
    if "目的" in title:
        return section_content(sections, "文档目的", "目的", used_keys=used_keys)
    if "背景" in title:
        content, key = section_content(sections, "项目概况", "背景", used_keys=used_keys)
        return content or context["project"].get("background", ""), key
    if "范围" in title and "测试" not in title:
        content, key = section_content(sections, "测试范围", used_keys=used_keys)
        return content, key
    if "实施任务" in title or "人员安排" in title:
        return "", ""
    if "整体进度" in title or "测试进度" in title:
        return section_content(sections, "测试阶段安排", "测试规划", used_keys=used_keys)
    if "硬件环境" in title or "服务器环境" in title:
        return section_content(sections, "服务器环境", "测试环境与工具", used_keys=used_keys)
    if "软件测试环境" in title or "环境要求" in title:
        return section_content(sections, "测试环境与工具", "测试环境拓扑", "服务器环境", used_keys=used_keys)
    if "测试工具" in title:
        return section_content(sections, "测试工具", used_keys=used_keys)
    if "功能测试" in title:
        return section_content(sections, "功能测试", "测试用例设计", used_keys=used_keys)
    if "接口测试" in title:
        return section_content(sections, "接口测试", used_keys=used_keys)
    if "性能测试" in title:
        return section_content(sections, "性能测试", used_keys=used_keys)
    if "测试方法" in title:
        return section_content(sections, "测试流程", "测试用例设计", used_keys=used_keys)
    if "通过标准" in title or "质量目标" in title or "完成准则" in title:
        return section_content(sections, "测试完成准则", used_keys=used_keys)
    if "风险" in title:
        return section_content(sections, "风险管理", used_keys=used_keys)
    return "", ""


def generate_template_driven_markdown(context: dict, template: dict, title: str | None, input_content: str) -> str | None:
    headings = template_headings(template.get("file_path", ""))
    if not headings or not input_content.strip():
        return None
    project = context["project"]
    doc_title = title or f"{project['name']} {TEMPLATE_TYPE_LABELS.get(template['template_type'], template['name'])}"
    sections = parse_markdown_sections(input_content)
    lines = [f"# {doc_title}", ""]
    used_keys: set[str] = set()
    used_any = False
    for index, (level, heading) in enumerate(headings):
        has_child = False
        for next_level, _ in headings[index + 1 :]:
            if next_level <= level:
                break
            has_child = True
        marker = "#" * min(max(level, 1) + 1, 4)
        lines.append(f"{marker} {heading}")
        if has_child:
            content, source_key = "", ""
        else:
            content, source_key = mapped_template_section(heading, sections, context, used_keys)
        if content:
            content = ground_project_content(content, context, template["template_type"])
        elif not has_child:
            content = fallback_section_content(heading, context, template["template_type"], used_keys)
        if content:
            lines.append("")
            lines.append(content)
            if source_key:
                used_keys.add(source_key)
            used_any = True
        lines.append("")
    if not used_any:
        return None
    return "\n".join(lines).strip() + "\n"


def generate_rule_based_markdown(
    context: dict,
    template: dict,
    title: str | None = None,
    input_content: str = "",
) -> str:
    project = context["project"]
    template_label = TEMPLATE_TYPE_LABELS.get(template["template_type"], template["name"])
    doc_title = title or f"{project['name']} {template_label}"
    template_driven = generate_template_driven_markdown(context, template, title, input_content)
    if template_driven:
        return template_driven
    if template["template_type"] == "weekly":
        return f"""# {doc_title}

## 1. 本周项目概况

- 项目名称：{project['name']}
- 客户/业主单位：{project['customer']}
- 当前阶段：{project['phase']}
- 项目经理：{project['manager']}
- 预算使用率：{project['budget_usage']}%

## 2. 本周任务进展

{bullet_list(context['tasks'], ('title', 'status', 'owner', 'due_date'), '暂无任务进展。')}

## 3. 风险与阻塞

{bullet_list(context['risks'], ('title', 'impact', 'status', 'response'), '暂无开放风险。')}

## 4. 新增需求与范围变化

{bullet_list(context['demands'], ('title', 'status', 'scope_impact'), '暂无新增需求。')}

## 5. 下周计划

{user_input_section(input_content) if input_content else project['plan'] or '待补充下周计划。'}
"""
    if template["template_type"] == "meeting_minutes":
        return f"""# {doc_title}

## 1. 会议背景

- 项目名称：{project['name']}
- 客户/业主单位：{project['customer']}
- 当前阶段：{project['phase']}

## 2. 会议输入

{user_input_section(input_content)}

## 3. 待办事项

{bullet_list(context['tasks'], ('title', 'owner', 'due_date', 'status'), '暂无待办事项。')}

## 4. 风险与决策

{bullet_list(context['risks'], ('title', 'impact', 'response'), '暂无风险或决策记录。')}
"""
    if template["template_type"] in {"api_design", "database_design", "technical_solution", "deployment_plan"}:
        return f"""# {doc_title}

## 1. 项目与目标

- 项目名称：{project['name']}
- 客户/业主单位：{project['customer']}
- 当前阶段：{project['phase']}
- 计划周期：{project['start_date'] or project['pre_start_date']} 至 {project['end_date'] or project['acceptance']}

## 2. 新增输入内容

{user_input_section(input_content)}

## 3. 设计范围

{project['plan'] or '待补充设计范围和实施计划。'}

## 4. 相关任务

{bullet_list(context['tasks'], ('title', 'description', 'owner', 'due_date'), '暂无任务数据。')}

## 5. 风险与约束

{bullet_list(context['risks'], ('title', 'description', 'impact', 'response'), '暂无风险数据。')}
"""
    if template["template_type"] in {"test_plan", "acceptance_report"}:
        return f"""# {doc_title}

## 1. 项目概况

- 项目名称：{project['name']}
- 客户/业主单位：{project['customer']}
- 当前阶段：{project['phase']}

## 2. 测试/验收输入

{user_input_section(input_content)}

## 3. 范围与检查项

{bullet_list(context['tasks'], ('title', 'description', 'status'), '暂无检查项。')}

## 4. 风险与遗留问题

{bullet_list(context['risks'], ('title', 'impact', 'status', 'response'), '暂无遗留风险。')}

## 5. 结论

待项目经理根据测试或验收结果补充。
"""
    if template["template_type"] in {"operation_manual", "user_manual"}:
        return f"""# {doc_title}

## 1. 适用范围

本文档适用于 {project['name']} 的交付、使用和维护场景。

## 2. 用户输入内容

{user_input_section(input_content)}

## 3. 项目基础信息

- 客户/业主单位：{project['customer']}
- 项目地区：{project['region']}
- 项目经理：{project['manager']}

## 4. 操作/维护事项

{bullet_list(context['tasks'], ('title', 'description', 'owner'), '暂无操作或维护事项。')}

## 5. 注意事项

{bullet_list(context['risks'], ('title', 'impact', 'response'), '暂无注意事项。')}
"""
    return f"""# {doc_title}

## 1. 项目背景

{project['background'] or '待补充项目背景。'}

## 2. 项目基础信息

- 客户/业主单位：{project['customer']}
- 项目地区：{project['region']}
- 当前阶段：{project['phase']}
- 合同状态：{project['contract_status']}
- 项目经理：{project['manager']}
- 计划周期：{project['start_date'] or project['pre_start_date']} 至 {project['end_date'] or project['acceptance']}

## 3. 功能需求与任务清单

### 3.1 新增输入内容

{user_input_section(input_content)}

### 3.2 已沉淀任务清单

{bullet_list(context['tasks'], ('title', 'description', 'owner', 'due_date'), '暂无任务数据。')}

## 4. 风险与待确认事项

{bullet_list(context['risks'], ('title', 'description', 'impact', 'response'), '暂无风险数据。')}

## 5. 需求池

{bullet_list(context['demands'], ('title', 'description', 'scope_impact'), '暂无新增需求。')}

## 6. 里程碑计划

{bullet_list(context['milestones'], ('title', 'date', 'status'), '暂无里程碑。')}
"""


def strip_empty_placeholder_lines(markdown: str) -> str:
    replacements = {
        "待补充。": "本节内容需结合项目最新资料确认后持续维护。",
        "暂无新增输入内容。本稿主要根据当前项目库上下文生成。": "本稿已根据当前项目库上下文生成，后续可结合会议纪要、需求清单或评审结论继续完善。",
    }
    cleaned = markdown
    for source, target in replacements.items():
        cleaned = cleaned.replace(source, target)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip() + "\n"


def remove_duplicate_tables(markdown: str) -> str:
    lines = markdown.splitlines()
    output: list[str] = []
    seen_tables: set[str] = set()
    index = 0
    while index < len(lines):
        table = parse_markdown_table(lines, index)
        if table:
            rows, next_index = table
            signature = "\n".join("|".join(row) for row in rows)
            if signature not in seen_tables:
                output.extend(lines[index:next_index])
                seen_tables.add(signature)
            index = next_index
            continue
        output.append(lines[index])
        index += 1
    return "\n".join(output)


def improve_document_markdown(markdown: str, context: dict, template: dict) -> str:
    project = context["project"]
    improved = ground_project_content(markdown, context, template.get("template_type", ""))
    if project.get("name"):
        lines = improved.splitlines()
        if lines and lines[0].startswith("# ") and not project["name"] in lines[0]:
            # Keep an explicit user title, but remove obvious stale project-title fragments from generated content.
            body = "\n".join(lines[1:])
            body = re.sub(r"项目名称[:：][^\n]+", f"项目名称：{project['name']}", body)
            lines = [lines[0], body]
            improved = "\n".join(lines)
    improved = remove_duplicate_tables(improved)
    improved = strip_empty_placeholder_lines(improved)
    return improved


def generate_markdown(
    context: dict,
    template: dict,
    title: str | None = None,
    input_content: str = "",
) -> str:
    draft = generate_rule_based_markdown(context, template, title, input_content)
    project = context["project"]
    doc_title = title or f"{project['name']} {TEMPLATE_TYPE_LABELS.get(template['template_type'], template['name'])}"
    try:
        from app.services.ai_adapter import generate_document_with_configured_model

        model_markdown = generate_document_with_configured_model(context, template, doc_title, input_content, draft)
    except Exception:
        model_markdown = None
    return improve_document_markdown(model_markdown or draft, context, template)


def is_markdown_table_separator(line: str) -> bool:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)


def parse_markdown_table(lines: list[str], start: int) -> tuple[list[list[str]], int] | None:
    if start + 1 >= len(lines) or "|" not in lines[start] or not is_markdown_table_separator(lines[start + 1]):
        return None
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and "|" in lines[index]:
        if index != start + 1:
            rows.append([cell.strip().replace("\\|", "|") for cell in lines[index].strip().strip("|").split("|")])
        index += 1
    return rows, index


def markdown_to_html(markdown: str) -> str:
    html_lines = []
    in_list = False
    lines = markdown.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index]
        table = parse_markdown_table(lines, index)
        if table:
            if in_list:
                html_lines.append("</ul>")
                in_list = False
            rows, next_index = table
            html_lines.append("<table>")
            for row_index, row in enumerate(rows):
                tag = "th" if row_index == 0 else "td"
                html_lines.append(
                    "<tr>"
                    + "".join(
                        f"<{tag}>{'<br>'.join(escape(part) for part in cell.split('<br>'))}</{tag}>"
                        for cell in row
                    )
                    + "</tr>"
                )
            html_lines.append("</table>")
            index = next_index
            continue
        if line.startswith("# "):
            if in_list:
                html_lines.append("</ul>")
                in_list = False
            html_lines.append(f"<h1>{escape(line[2:])}</h1>")
        elif line.startswith("## "):
            if in_list:
                html_lines.append("</ul>")
                in_list = False
            html_lines.append(f"<h2>{escape(line[3:])}</h2>")
        elif line.startswith("### "):
            if in_list:
                html_lines.append("</ul>")
                in_list = False
            html_lines.append(f"<h3>{escape(line[4:])}</h3>")
        elif line.startswith("- "):
            if not in_list:
                html_lines.append("<ul>")
                in_list = True
            html_lines.append(f"<li>{escape(line[2:])}</li>")
        elif line.strip():
            if in_list:
                html_lines.append("</ul>")
                in_list = False
            html_lines.append(f"<p>{escape(line)}</p>")
        index += 1
    if in_list:
        html_lines.append("</ul>")
    return "\n".join(html_lines)


def safe_filename(text: str) -> str:
    cleaned = re.sub(r"[^0-9A-Za-z\u4e00-\u9fa5_-]+", "-", text).strip("-")
    return cleaned or "document"


def iter_document_blocks(document: DocumentObject):
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def paragraph_heading_level(paragraph: Paragraph) -> int:
    style_name = paragraph.style.name if paragraph.style else ""
    match = re.search(r"Heading\s+(\d+)|标题\s*(\d+)", style_name, re.I)
    return int(match.group(1) or match.group(2) or 0) if match else 0


def markdown_escape_cell(value: str) -> str:
    lines = [re.sub(r"[ \t]+", " ", line.strip()) for line in value.strip().splitlines()]
    return "<br>".join(line for line in lines if line).replace("|", "\\|")


def table_to_markdown(table: Table) -> list[str]:
    rows: list[list[str]] = []
    max_cols = 0
    for row in table.rows:
        cells = [markdown_escape_cell(cell.text) for cell in row.cells]
        if any(cells):
            rows.append(cells)
            max_cols = max(max_cols, len(cells))
    if not rows:
        return []
    padded = [row + [""] * (max_cols - len(row)) for row in rows]
    header = padded[0]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join("---" for _ in header) + " |",
    ]
    for row in padded[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return lines


def should_skip_paragraph(paragraph: Paragraph, in_toc: bool) -> bool:
    text = paragraph.text.strip()
    style_name = paragraph.style.name if paragraph.style else ""
    if not text:
        return True
    if style_name.lower().startswith("toc") or "目录" in style_name or re.fullmatch(r"\d+(?:\.\d+)*\s+.+\t\d+", text):
        return True
    if in_toc and not paragraph_heading_level(paragraph):
        return True
    return False


def extract_docx_text(file_obj) -> str:
    document = Document(file_obj)
    lines: list[str] = []
    in_toc = False
    for block in iter_document_blocks(document):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            style_name = block.style.name if block.style else ""
            if text in {"目录", "目  录"} or "TOC 标题" in style_name:
                in_toc = True
                continue
            level = paragraph_heading_level(block)
            if in_toc and level:
                in_toc = False
            if should_skip_paragraph(block, in_toc):
                continue
            if level:
                lines.append(f"{'#' * min(level, 4)} {strip_heading_number(text)}")
            else:
                lines.append(text)
            lines.append("")
        else:
            table_lines = table_to_markdown(block)
            if table_lines:
                lines.extend(table_lines)
                lines.append("")
    return "\n".join(lines).strip()


def clear_document_body(document: DocumentObject) -> None:
    body = document.element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    paragraph.text = text


def replace_template_tokens(text: str, context: dict | None, title: str) -> str:
    if not context:
        return text.replace("XXX系统", title).replace("xxx系统", title).replace("Xxxx", "")
    project = context["project"]
    project_name = project.get("name") or title
    customer = project.get("customer") or "项目组"
    replacements = {
        "XXX系统": project_name,
        "xxx系统": project_name,
        "XXX": project_name,
        "Xxxx": customer,
        "xxxx": customer,
    }
    result = text
    for source, target in replacements.items():
        result = result.replace(source, str(target))
    return result


def markdown_section_parts(markdown: str) -> dict[str, dict[str, list]]:
    sections: dict[str, dict[str, list]] = {}
    current = ""
    lines = markdown.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index]
        level = markdown_heading_level(line)
        if level:
            current = normalize_heading(strip_heading_number(line[level + 1 :].strip()))
            sections.setdefault(current, {"paragraphs": [], "tables": []})
            index += 1
            continue
        table = parse_markdown_table(lines, index)
        if table and current:
            rows, next_index = table
            sections.setdefault(current, {"paragraphs": [], "tables": []})["tables"].append(rows)
            index = next_index
            continue
        if current and line.strip() and not line.strip().startswith(">"):
            sections.setdefault(current, {"paragraphs": [], "tables": []})["paragraphs"].append(line.strip())
        index += 1
    return sections


def first_section_text(sections: dict[str, dict[str, list]], *titles: str) -> str:
    for title in titles:
        section = sections.get(normalize_heading(strip_heading_number(title)))
        if section:
            paragraphs = [item for item in section["paragraphs"] if item and not item.startswith("- ")]
            if paragraphs:
                return "\n".join(paragraphs)
    return ""


def first_section_table(sections: dict[str, dict[str, list]], *titles: str) -> list[list[str]]:
    for title in titles:
        section = sections.get(normalize_heading(strip_heading_number(title)))
        if section and section["tables"]:
            return section["tables"][0]
    return []


def normalized_row_signature(row: list[str]) -> str:
    return "|".join(normalize_heading(cell) for cell in row)


def markdown_table_matching(
    sections: dict[str, dict[str, list]],
    required_headers: tuple[str, ...],
    *titles: str,
) -> list[list[str]]:
    required = [normalize_heading(header) for header in required_headers]
    candidates: list[list[list[str]]] = []
    for title in titles:
        section = sections.get(normalize_heading(strip_heading_number(title)))
        if section:
            candidates.extend(section["tables"])
    if not candidates:
        for section in sections.values():
            candidates.extend(section["tables"])
    for rows in candidates:
        if not rows:
            continue
        header = normalized_row_signature(rows[0])
        if all(item in header for item in required):
            return rows
    return []


def split_date_range(value: str) -> tuple[str, str, str]:
    text = value.replace("<br>", " ").replace("\n", " ").strip()
    date_pattern = r"\d{4}[./-]\d{1,2}[./-]\d{1,2}"
    dates = re.findall(date_pattern, text)
    stage = re.sub(rf"[（(]?\s*{date_pattern}\s*[~\-至到]\s*{date_pattern}\s*[）)]?", "", text).strip()
    stage = re.sub(r"\s{2,}", " ", stage) or text
    if len(dates) >= 2:
        return stage, dates[0].replace(".", "-").replace("/", "-"), dates[1].replace(".", "-").replace("/", "-")
    if len(dates) == 1:
        return stage, dates[0].replace(".", "-").replace("/", "-"), "待确认"
    return stage, "待确认", "待确认"


def schedule_rows_from_markdown(markdown_rows: list[list[str]], context: dict) -> list[list[str]]:
    if not markdown_rows:
        return []
    header = [normalize_heading(cell) for cell in markdown_rows[0]]
    tester = context["project"].get("delivery") or context["project"].get("manager") or "待确认"
    if "开始时间" in "".join(header) and "完成时间" in "".join(header):
        return markdown_rows
    rows = [["测试阶段", "开始时间", "完成时间", "测试人员"]]
    for row in markdown_rows[1:]:
        if not row or not any(cell.strip() for cell in row):
            continue
        stage, start, end = split_date_range(row[0])
        rows.append([stage, start, end, tester])
    return rows if len(rows) > 1 else []


def align_markdown_rows_to_template(markdown_rows: list[list[str]], template_header: str) -> list[list[str]]:
    if not markdown_rows:
        return []
    if "软件需求" in template_header and "用途" in template_header:
        return markdown_rows
    if "工具" in template_header:
        return markdown_rows
    if "机型配置" in template_header or "操作系统" in template_header:
        return markdown_rows
    return markdown_rows


def clear_table_body(table: Table) -> None:
    for row in table.rows[1:]:
        for cell in row.cells:
            cell.text = ""


def fill_table_preserve_shape(table: Table, rows: list[list[str]], keep_header: bool = True) -> None:
    if not rows:
        return
    start_row = 1 if keep_header else 0
    for row_index in range(start_row, len(table.rows)):
        source_index = row_index - start_row + (1 if keep_header else 0)
        source = rows[source_index] if source_index < len(rows) else []
        for cell_index, cell in enumerate(table.rows[row_index].cells):
            cell.text = source[cell_index] if cell_index < len(source) else ""


def normalized_table_header(table: Table) -> str:
    if not table.rows:
        return ""
    return "|".join(normalize_heading(cell.text) for cell in table.rows[0].cells)


def scope_rows(context: dict) -> list[list[str]]:
    items = context["demands"][:3] or context["tasks"][:3]
    rows = [["模块名称", "功能名称"]]
    if not items:
        rows += [["项目基础功能", "以需求清单和评审结论为准"], ["接口与集成", "按接口文档开展联调验证"], ["报表与查询", "按业务查询和统计口径验证"]]
        return rows
    for item in items:
        rows.append([str(item.get("title") or "待确认模块"), str(item.get("description") or item.get("scope_impact") or "功能点待确认")])
    return rows


def reference_doc_rows() -> list[list[str]]:
    return [
        ["文档 / （版本/日期）", "已创建或可用", "已被接收或已经过复审", "作者或来源", "备注"],
        ["需求规格说明书/需求清单", "是□　否□", "是□　否□", "项目组", "明确测试范围和验收口径"],
        ["原型图/界面设计稿", "是□　否□", "是□　否□", "产品/项目组", "支撑页面功能和交互验证"],
        ["接口设计文档", "是□　否□", "是□　否□", "研发/项目组", "支撑接口测试和联调验证"],
        ["部署实施文档", "是□　否□", "是□　否□", "交付/实施组", "支撑环境准备和版本部署检查"],
        ["数据库设计文档", "是□　否□", "是□　否□", "研发/项目组", "支撑数据一致性验证"],
        ["历史缺陷及变更记录", "是□　否□", "是□　否□", "测试/项目组", "支撑回归测试范围确认"],
        ["验收标准或客户确认记录", "是□　否□", "是□　否□", "客户/项目组", "支撑准出判断"],
        ["其他资料", "是□　否□", "是□　否□", "项目组", "按实际情况补充"],
    ]


def submission_doc_rows() -> list[list[str]]:
    return [
        ["文档说明", "作者", "文档位置"],
        ["测试计划/测试方案", "测试负责人", "项目文档库"],
        ["测试用例与执行记录", "测试负责人", "项目文档库"],
        ["缺陷清单、测试报告、验收问题清单", "测试/项目负责人", "项目文档库"],
    ]


def schedule_rows(context: dict) -> list[list[str]]:
    project = context["project"]
    tester = project.get("delivery") or project.get("manager") or "待确认"
    start, _ = project_period(context)
    stages = [
        ("测试准备", 0, 3),
        ("接口测试", 3, 5),
        ("系统测试", 8, 7),
        ("回归测试", 15, 3),
        ("测试总结", 18, 2),
    ]
    rows = [["测试阶段", "开始时间", "完成时间", "测试人员"]]
    for name, offset, days in stages:
        begin = start + timedelta(days=offset) if start else None
        end = begin + timedelta(days=days - 1) if begin else None
        rows.append([name, begin.isoformat() if begin else "待确认", end.isoformat() if end else "待确认", tester])
    return rows


def task_schedule_rows(context: dict) -> list[list[str]]:
    project = context["project"]
    tester = project.get("delivery") or project.get("manager") or "待确认"
    rows = [["测试功能点", "开始时间", "完成时间", "测试人员", "说明"]]
    tasks = context["tasks"][:3]
    if tasks:
        for task in tasks:
            rows.append([str(task.get("title") or "待确认"), str(task.get("start_date") or "待确认"), str(task.get("due_date") or "待确认"), str(task.get("owner") or tester), str(task.get("description") or "按测试计划执行")])
    else:
        rows += [
            ["功能测试", "待确认", "待确认", tester, "覆盖核心业务流程和异常场景"],
            ["接口测试", "待确认", "待确认", tester, "覆盖接口鉴权、参数校验和数据一致性"],
            ["回归测试", "待确认", "待确认", tester, "覆盖已修复缺陷和高风险流程"],
        ]
    return rows


def hardware_rows(context: dict) -> list[list[str]]:
    return [
        ["机型（配置）", "操作系统", "用途及特殊说明", "软件及版本", "预计空间"],
        ["应用服务器", "以部署方案为准", "部署被测系统服务", "应用服务组件", "按实际环境确认"],
        ["数据库服务器", "以部署方案为准", "数据库服务与数据验证", "数据库版本以部署文档为准", "按实际数据量确认"],
        ["测试客户端", "Windows/macOS", "测试执行、接口调试和问题复现", "浏览器/Postman/JMeter", "按测试工具需要"],
        ["网络环境", "客户内网/VPN", "访问被测系统和第三方接口", "网络白名单/证书", "按现场策略确认"],
    ]


def software_rows() -> list[list[str]]:
    return [
        ["软件需求", "用途"],
        ["操作系统、数据库、中间件版本以部署实施文档为准", "保证测试环境与交付环境一致"],
        ["浏览器、接口测试工具、缺陷管理工具", "支撑功能、接口、回归和缺陷闭环验证"],
    ]


def tool_rows() -> list[list[str]]:
    return [
        ["用途", "工具", "生产厂商/自产", "版本"],
        ["接口测试", "Postman/JMeter", "开源/第三方", "以项目环境为准"],
        ["性能测试", "JMeter", "开源", "以项目环境为准"],
        ["缺陷管理", "项目缺陷台账/禅道/飞书表格", "项目组", "以项目要求为准"],
        ["日志与监控", "系统日志/服务器监控工具", "项目组", "以部署环境为准"],
    ]


def method_rows(title: str) -> list[list[str]]:
    detail = test_method_text(title, {})
    return [
        ["测试目标", detail.splitlines()[0] if detail else f"验证{title}相关目标是否满足要求。"],
        ["测试范围：", f"{title}涉及的核心流程、异常分支和边界条件。"],
        ["开始标准：", "需求范围确认，测试环境、账号、数据和版本准备完成。"],
        ["完成标准：", "测试用例执行完成，阻塞和严重缺陷关闭，遗留问题完成确认。"],
        ["测试重点和优先级：", "优先覆盖核心业务流程、高风险模块和客户关注场景。"],
        ["需考虑的特殊事项：", "记录测试数据、操作步骤、实际结果、缺陷编号和处理结论。"],
        ["备注：", "具体用例以测试用例附件或项目文档库为准。"],
    ]


def quality_rows() -> list[list[str]]:
    return [
        ["测试质量目标", "确认者（如需说明）"],
        ["核心用例通过率达到 100%", "项目经理/测试负责人"],
        ["总体用例通过率不低于 95%", "测试负责人"],
        ["阻塞和严重缺陷为 0 个未关闭", "测试负责人/开发负责人"],
        ["一般缺陷经评审可遗留并明确计划", "项目经理"],
        ["遗留问题形成书面记录并纳入后续跟踪", "项目经理/客户代表"],
        ["测试报告输出并完成评审", "项目经理/测试负责人"],
        ["验收问题清单完成确认", "客户代表/项目经理"],
    ]


def defect_rows() -> list[list[str]]:
    return [
        ["问题严重程度", "描述"],
        ["致命", "系统不可用、核心流程完全阻断或造成严重数据错误。"],
        ["严重", "重要功能不可用、权限异常或影响关键业务办理。"],
        ["一般", "局部功能异常、提示不准确或影响操作效率。"],
        ["轻微", "文案、样式、兼容性等轻微问题，可结合版本节奏处理。"],
    ]


def usable_width_twips(document: DocumentObject) -> int:
    section = document.sections[-1]
    return int((section.page_width - section.left_margin - section.right_margin) / 635)


def set_table_layout(table: Table, width_twips: int) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_twips))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def table_column_widths(rows: list[list[str]], width_twips: int) -> list[int]:
    max_cols = max(len(row) for row in rows)
    scores: list[float] = []
    for column in range(max_cols):
        values = [row[column] if column < len(row) else "" for row in rows]
        score = max((len(value.replace("<br>", "")) for value in values), default=0)
        scores.append(max(8, min(score, 42)) ** 0.75)
    total = sum(scores) or max_cols
    widths = [max(900, int(width_twips * score / total)) for score in scores]
    overflow = sum(widths) - width_twips
    if overflow > 0:
        widest = max(range(len(widths)), key=lambda item: widths[item])
        widths[widest] = max(900, widths[widest] - overflow)
    return widths


def format_cell(cell, value: str, width_twips: int, is_header: bool) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_width(cell, width_twips)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    for index, part in enumerate(value.split("<br>")):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(part)
        run.bold = is_header
        run.font.size = Pt(9)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_markdown_table(document: DocumentObject, rows: list[list[str]]) -> None:
    if not rows:
        return
    max_cols = max(len(row) for row in rows)
    width_twips = usable_width_twips(document)
    widths = table_column_widths(rows, width_twips)
    table = document.add_table(rows=0, cols=max_cols)
    table.style = "Table Grid"
    set_table_layout(table, width_twips)
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for cell_index in range(max_cols):
            value = row[cell_index] if cell_index < len(row) else ""
            format_cell(cells[cell_index], value, widths[cell_index], row_index == 0)


def add_bullet_paragraph(document: DocumentObject, text: str) -> None:
    try:
        document.add_paragraph(text, style="List Bullet")
    except KeyError:
        document.add_paragraph(f"• {text}")


def markdown_title(markdown: str, fallback: str = "") -> str:
    for line in markdown.splitlines():
        if line.startswith("# "):
            return line[2:].strip()
    return fallback


def table_text_signature(table: Table) -> str:
    return "|".join(normalize_heading(cell.text) for row in table.rows[:2] for cell in row.cells)


def set_table_cell_text(table: Table, row_index: int, cell_index: int, text: str) -> None:
    if row_index < len(table.rows) and cell_index < len(table.rows[row_index].cells):
        table.rows[row_index].cells[cell_index].text = text


def fill_cover_or_revision_table(table: Table, context: dict, title: str) -> bool:
    signature = table_text_signature(table)
    project = context["project"]
    project_name = project.get("name") or title
    customer = project.get("customer") or "项目组"
    today = date.today().isoformat()
    filled = False
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            text = normalize_heading(cell.text)
            if "编制单位" in text and cell_index + 1 < len(row.cells):
                row.cells[cell_index + 1].text = customer
                filled = True
            elif "编制日期" in text and cell_index + 1 < len(row.cells):
                row.cells[cell_index + 1].text = today
                filled = True
            elif "系统名称" in text and cell_index + 1 < len(row.cells):
                row.cells[cell_index + 1].text = project_name
                filled = True
            elif "项目名称" in text and cell_index + 1 < len(row.cells):
                row.cells[cell_index + 1].text = project_name
                filled = True
            elif "版本" in text and cell_index + 1 < len(row.cells) and not row.cells[cell_index + 1].text.strip():
                row.cells[cell_index + 1].text = "V1.0"
                filled = True
            else:
                cell.text = replace_template_tokens(cell.text, context, title)
    if not filled and any(keyword in signature for keyword in ("编制单位", "修订记录", "版本号", "文档编号")):
        filled = True
    return filled


def fill_template_table(
    table: Table,
    current_heading: str,
    context: dict,
    title: str,
    sections: dict[str, dict[str, list]],
) -> bool:
    header = normalized_table_header(table)
    signature = table_text_signature(table)
    heading = strip_heading_number(current_heading)
    if fill_cover_or_revision_table(table, context, title):
        return True
    if "模块名称" in header and "功能名称" in header:
        markdown_rows = markdown_table_matching(sections, ("模块名称", "功能名称"), heading, "范围", "测试范围", "范围与检查项")
        fill_table_preserve_shape(table, markdown_rows or scope_rows(context), keep_header=True)
        return True
    if ("文档版本日期" in header or "文档版本" in header) and ("已创建或可用" in header or "已被接收" in header):
        markdown_rows = markdown_table_matching(sections, ("文档", "已创建"), heading, "测试参考文档", "参考文档")
        fill_table_preserve_shape(table, markdown_rows or reference_doc_rows(), keep_header=True)
        return True
    if "文档说明" in header and "作者" in header and "文档位置" in header:
        markdown_rows = markdown_table_matching(sections, ("文档", "作者"), heading, "测试提交文档", "提交文档", "交付文档")
        fill_table_preserve_shape(table, markdown_rows or submission_doc_rows(), keep_header=True)
        return True
    if "测试阶段" in header and "开始时间" in header and "完成时间" in header and "测试人员" in header:
        markdown_rows = markdown_table_matching(sections, ("测试阶段",), heading, "整体进度", "测试进度", "测试阶段安排")
        fill_table_preserve_shape(table, schedule_rows_from_markdown(markdown_rows, context) or schedule_rows(context), keep_header=True)
        return True
    if "测试功能点" in header and "开始时间" in header and "完成时间" in header:
        markdown_rows = markdown_table_matching(sections, ("测试功能点",), heading, "实施任务和时间人员安排", "任务安排")
        fill_table_preserve_shape(table, align_markdown_rows_to_template(markdown_rows, header) or task_schedule_rows(context), keep_header=True)
        return True
    if "机型配置" in header and "操作系统" in header:
        markdown_rows = markdown_table_matching(sections, ("操作系统",), heading, "硬件环境", "服务器环境", "环境要求")
        fill_table_preserve_shape(table, align_markdown_rows_to_template(markdown_rows, header) or hardware_rows(context), keep_header=True)
        return True
    if "软件需求" in header and "用途" in header:
        markdown_rows = markdown_table_matching(sections, ("软件", "用途"), heading, "软件测试环境", "软件环境", "环境要求")
        fill_table_preserve_shape(table, align_markdown_rows_to_template(markdown_rows, header) or software_rows(), keep_header=True)
        return True
    if "工具" in header and "生产厂商自产" in header:
        markdown_rows = markdown_table_matching(sections, ("工具",), heading, "测试工具", "环境要求")
        fill_table_preserve_shape(table, align_markdown_rows_to_template(markdown_rows, header) or tool_rows(), keep_header=True)
        return True
    if "测试目标" in signature and any(keyword in heading for keyword in ("功能测试", "接口测试", "集成测试", "用户界面测试", "性能测试", "回归测试", "负载测试", "安全测试")):
        fill_table_preserve_shape(table, method_rows(heading), keep_header=False)
        return True
    if "测试质量目标" in header or ("测试质量目标" in signature and "确认者" in signature):
        fill_table_preserve_shape(table, quality_rows(), keep_header=True)
        return True
    if "问题严重程度" in header and "描述" in header:
        fill_table_preserve_shape(table, defect_rows(), keep_header=True)
        return True
    return False


def template_section_text(heading: str, sections: dict[str, dict[str, list]], context: dict, template_type: str) -> str:
    normalized = normalize_heading(strip_heading_number(heading))
    direct = sections.get(normalized, {})
    paragraphs = direct.get("paragraphs", []) if direct else []
    if paragraphs:
        return "\n".join(str(item) for item in paragraphs if item and not str(item).startswith("- "))
    if "目的" in heading:
        return first_section_text(sections, "文档目的", "目的") or fallback_section_content("目的", context, template_type, set())
    if "背景" in heading:
        return first_section_text(sections, "项目概况", "背景") or fallback_section_content("背景", context, template_type, set())
    if "范围" in heading and "测试" not in heading:
        return first_section_text(sections, "测试范围", "范围") or fallback_section_content("范围", context, template_type, set())
    if "风险" in heading:
        return first_section_text(sections, "风险管理", "风险和应急") or fallback_section_content("风险", context, template_type, set())
    if "通过标准" in heading or "完成准则" in heading:
        return first_section_text(sections, "测试完成准则", "测试项通过标准") or fallback_section_content("测试完成准则", context, template_type, set())
    return ""


def export_markdown_to_template_docx(
    markdown: str,
    output_path: Path,
    template_path: str,
    context: dict,
    template: dict | None = None,
) -> None:
    document = Document(template_path)
    title = markdown_title(markdown, context["project"].get("name") or "")
    sections = markdown_section_parts(markdown)
    template_type = (template or {}).get("template_type", "test_plan")
    current_heading = ""
    filled_paragraph_sections: set[str] = set()
    for block in iter_document_blocks(document):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            level = paragraph_heading_level(block)
            if level:
                current_heading = strip_heading_number(text)
                set_paragraph_text(block, replace_template_tokens(block.text, context, title))
                continue
            if not text:
                continue
            section_key = normalize_heading(current_heading)
            replacement = template_section_text(current_heading, sections, context, template_type) if current_heading else ""
            if replacement and section_key not in filled_paragraph_sections:
                set_paragraph_text(block, replace_template_tokens(replacement, context, title))
                filled_paragraph_sections.add(section_key)
            else:
                set_paragraph_text(block, replace_template_tokens(block.text, context, title))
        else:
            if fill_template_table(block, current_heading, context, title, sections):
                continue
            for row in block.rows:
                for cell in row.cells:
                    cell.text = replace_template_tokens(cell.text, context, title)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def export_markdown_to_docx(
    markdown: str,
    output_path: Path,
    template_path: str = "",
    context: dict | None = None,
    template: dict | None = None,
) -> None:
    if (
        context
        and template_path
        and Path(template_path).suffix.lower() == ".docx"
        and Path(template_path).exists()
        and Document(template_path).tables
    ):
        export_markdown_to_template_docx(markdown, output_path, template_path, context, template)
        return
    if template_path and Path(template_path).suffix.lower() == ".docx" and Path(template_path).exists():
        document = Document(template_path)
        clear_document_body(document)
    else:
        document = Document()
    lines = markdown.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index]
        table = parse_markdown_table(lines, index)
        if table:
            rows, next_index = table
            add_markdown_table(document, rows)
            index = next_index
            continue
        if line.startswith("# "):
            document.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            document.add_heading(line[4:], level=3)
        elif line.startswith("#### "):
            document.add_heading(line[5:], level=4)
        elif line.startswith("- "):
            add_bullet_paragraph(document, line[2:])
        elif line.strip():
            document.add_paragraph(line)
        index += 1
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
